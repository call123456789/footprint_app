from PyQt5.QtWidgets import QDialog, QVBoxLayout, QTextEdit, QPushButton, QLabel, QMessageBox, QApplication, QProgressDialog
from PyQt5.QtCore import Qt, QTimer, pyqtSignal, QThread
from llm import LLM  # 确保LLM类在llm.py中定义
from docx import Document  # 确保安装了python-docx库

class PlanDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("生成旅行计划")
        self.setFixedSize(600, 400)

        layout = QVBoxLayout()

        # 提示文字
        self.label = QLabel("在此输入你对本次国内旅行的要求，如参观景点，交通方式，时间安排，住宿方式等")
        self.label.setWordWrap(True)  # 允许自动换行
        self.label.setAlignment(Qt.AlignCenter)
        layout.addWidget(self.label)

        # 输入框
        self.text_edit = QTextEdit()
        layout.addWidget(self.text_edit)

        # 确认按钮
        self.confirm_button = QPushButton("确定")
        self.confirm_button.clicked.connect(self.generate_plan)
        layout.addWidget(self.confirm_button)

        self.setLayout(layout)

    def generate_plan(self):
        """生成旅行计划"""
        requirements = self.text_edit.toPlainText()
        if not requirements.strip():
            QMessageBox.warning(self, "提示", "请输入旅行要求")
            return

        # 显示自定义进度对话框
        self.progress_dialog = QDialog(self)
        self.progress_dialog.setWindowTitle("提示")
        self.progress_dialog.setFixedSize(300, 150)

        progress_layout = QVBoxLayout()
        self.progress_label = QLabel("正在生成中...")
        self.progress_label.setAlignment(Qt.AlignCenter)
        progress_layout.addWidget(self.progress_label)

        self.progress_dialog.setLayout(progress_layout)
        self.progress_dialog.setModal(True)  # 设置为模态对话框
        self.progress_dialog.show()
        # 调用API生成计划
        prompt = f"根据以下要求生成旅行计划：\n{requirements}"
        self.ai_thread = self.AIWorker(
            prompt=prompt,
            llm=LLM()
        )
        self.ai_thread.ai_response.connect(self._generate_plan_async)
        self.ai_thread.start()


    class AIWorker(QThread):
        ai_response = pyqtSignal(str)  # 信号参数：(上方内容, 下方内容)

        def __init__(self, prompt,llm):
            super().__init__()
            self.prompt = prompt
            self.llm = llm

        def run(self):
            response = self.llm.response(
                    self.prompt
                )
            self.ai_response.emit(response)

    def _generate_plan_async(self, plan_text):
        """异步生成计划"""

        # 将计划保存为Word文档
        document = Document()
        document.add_heading("旅行计划", level=1)
        document.add_paragraph(plan_text)
        document.save("旅行计划.docx")

        # 隐藏进度对话框
        self.progress_dialog.close()

        QMessageBox.information(self, "提示", "旅行计划已生成并保存为Word文档")
        self.close()